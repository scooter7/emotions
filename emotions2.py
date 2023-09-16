import streamlit as st
import cv2
import numpy as np
from deepface import DeepFace
from docx import Document
from docx.shared import Inches
import base64
import os

st.title("DeepFace Analysis")
image_file = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png"])

last_analyzed_image = None
last_analysis = {}

def analyze_image(image_file):
    image_data = np.frombuffer(image_file.read(), np.uint8)
    if len(image_data) == 0:
        st.write("Could not read image data.")
        return
    
    image = cv2.imdecode(image_data, -1)
    if image is None:
        st.write("Could not decode image.")
        return

    image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    if image_rgb is None:
        st.write("Could not convert image to RGB.")
        return

    predictions = DeepFace.analyze(image_rgb, actions=['emotion'])

    if not isinstance(predictions, dict):
        st.write("Returned value:", predictions)
        return

    if 'dominant_emotion' not in predictions:
        st.write("DeepFace.analyze output:", predictions)
        return

    st.write("Emotion:", predictions['dominant_emotion'])

if image_file is not None:
    st.image(image_file, caption="Uploaded Image.", use_column_width=True)
    st.write("Select Analysis:")
    if st.button("Analyze Emotion"):
        analyze_image('emotion')
    if st.button("Analyze Age"):
        analyze_image('age')
    if st.button("Analyze Gender"):
        analyze_image('gender')
    if st.button("Analyze Race"):
        analyze_image('race')

def generate_word_file():
    global last_analyzed_image, last_analysis
    doc = Document()
    doc.add_heading('DeepFace Analysis Report', 0)
    if last_analyzed_image is not None:
        image_path = 'temp_image.jpg'
        cv2.imwrite(image_path, cv2.cvtColor(last_analyzed_image, cv2.COLOR_RGB2BGR))
        doc.add_picture(image_path, width=Inches(2.0))
        os.remove(image_path)
        for key, value in last_analysis.items():
            doc.add_paragraph(f'{key}: {value}')
    doc.save('report.docx')
    with open('report.docx', 'rb') as f:
        bytes = f.read()
        b64 = base64.b64encode(bytes).decode()
        href = f'<a href="data:file/docx;base64,{b64}" download="report.docx">Download Report</a>'
        st.markdown(href, unsafe_allow_html=True)

if st.button("Generate Report"):
    generate_word_file()
